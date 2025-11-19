#!/usr/bin/env python3
"""
Convert markdown to Word document using template
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path
import re

def convert_md_to_docx(md_path: str, template_path: str, output_path: str):
    """Convert markdown file to Word document using template"""
    
    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Load template
    doc = Document(template_path)
    
    # Clear existing content but keep structure
    # Find and replace title
    title_replaced = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if 'ENTER SERVICE NAME HERE' in text or 'G-Cloud' in text or not title_replaced:
            para.text = 'G-Cloud Proposal Automation Application (Azure Production)'
            for run in para.runs:
                run.font.size = Pt(24)
                run.font.bold = True
            title_replaced = True
            break
    
    # Remove all other paragraphs except title
    paragraphs_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        if i > 0:  # Keep first paragraph (title)
            paragraphs_to_remove.append(para)
    
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)
    
    # Parse markdown sections
    sections = content.split('### ')
    
    for section in sections[1:]:  # Skip first empty section
        lines = section.strip().split('\n')
        if not lines:
            continue
            
        heading = lines[0].strip()
        body_lines = lines[1:] if len(lines) > 1 else []
        
        # Add heading
        doc.add_paragraph(heading, style='Heading 2')
        
        # Process body
        current_list = None
        in_code_block = False
        
        for line in body_lines:
            line = line.rstrip()
            
            # Skip code block markers
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                # Add as monospace paragraph
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.name = 'Courier New'
                continue
            
            # Handle bullet points
            if line.strip().startswith('- '):
                text = line.strip()[2:].strip()
                doc.add_paragraph(text, style='List Bullet')
                current_list = True
            elif line.strip().startswith('|'):
                # Table row - skip for now
                continue
            elif line.strip() and not line.strip().startswith('#'):
                # Regular paragraph
                if current_list:
                    # End list
                    current_list = None
                doc.add_paragraph(line.strip())
            elif line.strip().startswith('#'):
                # Sub-heading
                level = len(line) - len(line.lstrip('#'))
                style = f'Heading {min(level + 2, 9)}'
                doc.add_paragraph(line.strip().lstrip('#').strip(), style=style)
    
    # Save
    doc.save(output_path)
    print(f"Document created: {output_path}")

if __name__ == '__main__':
    md_path = 'Developer_Guides/G-Cloud_Proposal_Automation_Application_Azure.md'
    template_path = 'docs/service_description_template.docx'
    output_path = 'Developer_Guides/G-Cloud_Proposal_Automation_Application_Azure.docx'
    
    convert_md_to_docx(md_path, template_path, output_path)

