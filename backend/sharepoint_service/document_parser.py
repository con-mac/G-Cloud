"""
Word document parser for extracting content from existing documents
"""

from pathlib import Path
from typing import Dict, List, Optional, Union
from docx import Document
from io import BytesIO
import logging
import re

logger = logging.getLogger(__name__)


def parse_service_description_document(doc_path: Union[Path, BytesIO, str]) -> Dict:
    """
    Parse a Service Description Word document and extract content.
    
    Args:
        doc_path: Path to Word document (.docx file)
        
    Returns:
        Dict with extracted content:
        {
            'title': str,
            'description': str,
            'features': List[str],
            'benefits': List[str],
            'service_definition': List[{'subtitle': str, 'content': str}]
        }
    """
    try:
        # Handle both Path objects and BytesIO (for S3)
        if isinstance(doc_path, BytesIO):
            doc = Document(doc_path)
        else:
            doc = Document(str(doc_path))
        
        result = {
            'title': '',
            'description': '',
            'features': [],
            'benefits': [],
            'service_definition': []
        }
        
        current_section = None
        current_content = []
        in_features = False
        in_benefits = False
        in_service_def = False
        current_subsection = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            style = para.style.name if para.style else ''
            
            # Title (Heading 1)
            if style.startswith('Heading 1') or (not result['title'] and len(text) < 100):
                if not result['title']:
                    result['title'] = text
                    continue
            
            # Short Service Description
            if 'Short Service Description' in text or 'short service description' in text.lower():
                current_section = 'description'
                continue
            
            # Key Service Features
            if 'Key Service Features' in text or 'key service features' in text.lower():
                current_section = 'features'
                in_features = True
                in_benefits = False
                in_service_def = False
                continue
            
            # Key Service Benefits
            if 'Key Service Benefits' in text or 'key service benefits' in text.lower():
                current_section = 'benefits'
                in_features = False
                in_benefits = True
                in_service_def = False
                continue
            
            # Service Definition (Heading 3)
            if style.startswith('Heading 3'):
                # Save previous subsection if exists
                if current_subsection:
                    result['service_definition'].append(current_subsection)
                # Start new subsection
                # Replace AI Security advisory with Lorem Ipsum
                subtitle = text
                if 'AI Security' in subtitle or 'advisory' in subtitle.lower() or '1.4.1' in subtitle:
                    subtitle = 'Lorem ipsum dolor sit amet'
                current_subsection = {'subtitle': subtitle, 'content': ''}
                in_features = False
                in_benefits = False
                in_service_def = True
                continue
            
            # Collect content based on current section
            if current_section == 'description' and not in_features and not in_benefits:
                if result['description']:
                    result['description'] += ' ' + text
                else:
                    result['description'] = text
            
            elif in_features:
                # Strip numbered prefixes (1., 2., etc.) from features
                # Numbers are for Word document formatting only - should not appear in form
                # The Word generator creates: Run 1 = "1. " (red), Run 2 = "text" (black)
                # When we read para.text, it combines both runs as "1. text"
                original_text = text.strip()
                stripped_text = re.sub(r'^\s*\d+[\.\)]?\s*', '', original_text)
                if original_text != stripped_text:
                    logger.debug(f"Stripped number from feature: '{original_text}' -> '{stripped_text}'")
                if stripped_text and not any(keyword in stripped_text.lower() for keyword in ['key service', 'short service']):
                    result['features'].append(stripped_text)
            
            elif in_benefits:
                # Strip numbered prefixes (1., 2., etc.) from benefits
                # Numbers are for Word document formatting only - should not appear in form
                # The Word generator creates: Run 1 = "1. " (red), Run 2 = "text" (black)
                # When we read para.text, it combines both runs as "1. text"
                original_text = text.strip()
                stripped_text = re.sub(r'^\s*\d+[\.\)]?\s*', '', original_text)
                if original_text != stripped_text:
                    logger.debug(f"Stripped number from benefit: '{original_text}' -> '{stripped_text}'")
                if stripped_text and not any(keyword in stripped_text.lower() for keyword in ['key service', 'short service']):
                    result['benefits'].append(stripped_text)
            
            elif in_service_def and current_subsection:
                # For service definition, preserve HTML format if possible
                # Check if paragraph has inline shapes (images) - for now just preserve text
                # Images will need more complex extraction logic
                if current_subsection['content']:
                    current_subsection['content'] += ' ' + text
                else:
                    current_subsection['content'] = text
        
        # Save last subsection if exists
        if current_subsection:
            result['service_definition'].append(current_subsection)
        
        # Clean up empty strings
        result['features'] = [f for f in result['features'] if f.strip()]
        result['benefits'] = [b for b in result['benefits'] if b.strip()]
        
        return result
        
    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        raise


def parse_service_description_document_from_bytes(doc_bytes: BytesIO) -> Dict:
    """
    Parse a Service Description Word document from bytes (for S3).
    
    Args:
        doc_bytes: BytesIO object containing Word document
        
    Returns:
        Dict with extracted content (same format as parse_service_description_document)
    """
    return parse_service_description_document(doc_bytes)


def read_document_content(service_name: str, doc_type: str, lot: str, gcloud_version: str = "14") -> Optional[Dict]:
    """
    Read and parse a document from SharePoint (local or S3).
    Works with both local (mock_sharepoint) and S3 storage.
    
    Args:
        service_name: Service name
        doc_type: Document type ("SERVICE DESC" or "Pricing Doc")
        lot: LOT number ("2" or "3")
        gcloud_version: GCloud version ("14" or "15")
        
    Returns:
        Parsed document content or None if not found
    """
    import os
    from io import BytesIO
    
    # Import SharePoint service abstraction (switches between local and S3)
    try:
        try:
            from sharepoint_service.sharepoint_service import get_document_path, USE_S3
        except ImportError:
            from app.sharepoint_service.sharepoint_service import get_document_path, USE_S3
    except ImportError as e:
        logger.error(f"Failed to import SharePoint service: {e}")
        return None
    
    doc_path = get_document_path(service_name, doc_type, lot, gcloud_version)
    
    if not doc_path:
        return None
    
    # Handle S3 storage
    if USE_S3:
        # doc_path is an S3 key (string)
        try:
            import boto3
            s3_client = boto3.client('s3')
            bucket_name = os.environ.get('SHAREPOINT_BUCKET_NAME', '')
            
            if not bucket_name:
                logger.error("SHAREPOINT_BUCKET_NAME not set")
                return None
            
            # Download document from S3 to memory
            response = s3_client.get_object(Bucket=bucket_name, Key=doc_path)
            doc_bytes = response['Body'].read()
            
            # Parse from bytes
            if doc_type == "SERVICE DESC":
                # Create a temporary file-like object from bytes
                doc_file = BytesIO(doc_bytes)
                return parse_service_description_document_from_bytes(doc_file)
            
            return None
        except Exception as e:
            logger.error(f"Error reading document from S3: {e}")
            return None
    
    # Handle local storage
    # doc_path is a Path object
    if not doc_path.exists():
        return None
    
    if doc_type == "SERVICE DESC":
        return parse_service_description_document(doc_path)
    
    # For other document types, we can extend later
    return None

