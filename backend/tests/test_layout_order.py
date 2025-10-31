from pathlib import Path
from zipfile import ZipFile
import os
from docx import Document
from app.services.document_generator import document_generator


def _make_minimal_template(path: Path):
    doc = Document()
    # Cover title placeholder
    doc.add_paragraph("Add Title")
    # Contents heading
    p = doc.add_paragraph("Contents")
    p.style = "Heading 2"
    # Some static content that should be cleared
    doc.add_paragraph("STATIC CONTENT BELOW TO BE CLEARED")
    doc.add_paragraph("Old text that must go")
    # About PA marker and heading (to be moved to the end)
    doc.add_paragraph("{{ABOUT_PA_START}}")
    ap = doc.add_paragraph("About PA")
    ap.style = "Heading 2"
    doc.add_paragraph("About PA body text")
    doc.save(path)


def test_order_and_about_pa_last(tmp_path, monkeypatch):
    # Create a minimal template with Contents + marker + About PA
    tpl_path = Path("/tmp/test_layout_template.docx")
    _make_minimal_template(tpl_path)
    monkeypatch.setenv("SERVICE_DESC_TEMPLATE_PATH", str(tpl_path))

    title = "My New Service"
    description = "Short description for testing"
    features = ["Feature one", "Feature two"]
    benefits = ["Benefit one", "Benefit two"]

    result = document_generator.generate_service_description(
        title=title,
        description=description,
        features=features,
        benefits=benefits,
        service_definition=[{"subtitle": "Sect A", "content": "<p>HTML block</p>"}],
    )

    docx_path = Path(result["word_path"])
    assert docx_path.exists()

    with ZipFile(docx_path, "r") as z:
        xml = z.read("word/document.xml").decode("utf-8")
        # Title must appear
        assert title in xml
        # About PA heading must appear
        assert "About PA" in xml
        # Service Definition heading appears
        assert "Service Definition" in xml

        # Order check: About PA content should come AFTER the inserted sections
        idx_title = xml.find(title)
        idx_sd = xml.find("Short Service Description")
        idx_features = xml.find("Key Service Features")
        idx_benefits = xml.find("Key Service Benefits")
        idx_about = xml.find("About PA")
        assert idx_title != -1 and idx_sd != -1 and idx_features != -1 and idx_benefits != -1 and idx_about != -1
        assert idx_about > max(idx_title, idx_sd, idx_features, idx_benefits)

        # Ensure marker was removed
        assert "{{ABOUT_PA_START}}" not in xml

