from pathlib import Path
from app.services.document_generator import document_generator


def test_numbering_resets_between_sections(tmp_path, monkeypatch):
    # Use a minimal template in /tmp
    tpl = tmp_path / "tpl.docx"
    from docx import Document
    d = Document()
    p = d.add_paragraph("Contents"); p.style = "Heading 2"
    d.add_paragraph("{{ABOUT_PA_START}}")
    ap = d.add_paragraph("About PA"); ap.style = "Heading 2"
    d.add_paragraph("About text")
    d.save(tpl)
    monkeypatch.setenv("SERVICE_DESC_TEMPLATE_PATH", str(tpl))

    title = "Svc"
    desc = "desc"
    feats = ["f1", "f2", "f3"]
    bens = ["b1", "b2"]
    result = document_generator.generate_service_description(
        title=title,
        description=desc,
        features=feats,
        benefits=bens,
        service_definition=None,
    )

    docx = Path(result["word_path"]).read_bytes()
    xml = None
    from zipfile import ZipFile
    from io import BytesIO
    with ZipFile(BytesIO(docx), 'r') as z:
        xml = z.read('word/document.xml').decode('utf-8')
    # After 'Key Service Features' there should be '1. '
    kf = xml.find('Key Service Features')
    kb = xml.find('Key Service Benefits')
    assert kf != -1 and kb != -1
    first_num_after_features = xml.find('1. ', kf)
    first_num_after_benefits = xml.find('1. ', kb)
    assert first_num_after_features != -1
    assert first_num_after_benefits != -1

