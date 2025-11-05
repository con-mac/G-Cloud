"""
Script to fix the service_description_template.docx by replacing AI Security content
with generic placeholders.
"""

from pathlib import Path
from docx import Document
import shutil
from datetime import datetime

def fix_template():
    """Replace AI Security content with generic placeholders"""
    
    template_path = Path(__file__).parent.parent / "templates" / "service_description_template.docx"
    backup_path = template_path.parent / f"service_description_template_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    print(f"Loading template: {template_path}")
    
    # DO NOT create backup - contains classified content
    # if template_path.exists():
    #     shutil.copy2(template_path, backup_path)
    #     print(f"Backup created: {backup_path}")
    
    doc = Document(str(template_path))
    
    # Replacements to make
    replacements = {
        # Title
        'AI Security': 'ENTER SERVICE NAME HERE',
        
        # Description - catch all AI Security description content
        'AI Security Advisory': '[SERVICE NAME]',
        'AI security': '[service area]',
        'AI Security': '[SERVICE NAME]',
        'enables teams to understand and use AI ethically and securely': '[generic service description placeholder]',
        'Secure Design and Development of AI helps': '[generic service description placeholder]',
        'AI products are secure-by-design': '[generic service description placeholder]',
        'use case-relevant': '[generic service description placeholder]',
        
        # Features - replace with generic placeholders
        'Map AI use across the enterprise': 'Feature placeholder 1',
        'Conduct AI security risk health check': 'Feature placeholder 2',
        'Help leaders navigate a complex and evolving landscape': 'Feature placeholder 3',
        'Application threat assessment, monitoring and maintenance': 'Feature placeholder 4',
        'AI Secure-by-design, secure development, Testing and Deployment': 'Feature placeholder 5',
        'Identify attacks using MITRE ATLAS, Microsoft Failure Modes, Google SAIF': 'Feature placeholder 6',
        'Validate your defence in depth to identify the best places to deploy AI defences.': 'Feature placeholder 7',
        'Specify AI specific requirements and architecture principles': 'Feature placeholder 8',
        'Develop security roadmap aligned with organisational and technological ambition': 'Feature placeholder 9',
        'Look across people, process, and technology to innovate your defences': 'Feature placeholder 10',
        
        # Benefits - replace with generic placeholders
        'Understand IT platforms hosting AI capabilities, enumerate systems at risk': 'Benefit placeholder 1',
        'Identify shadow AI and third-party risk. Assess mitigation best practice': 'Benefit placeholder 2',
        'Build security teams AI literacy to enshrine AI security governance': 'Benefit placeholder 3',
        'AI literacy': '[service area] literacy',
        'Implement AI-aware data governance and protection': 'Benefit placeholder 4',
        'Use cutting edge approaches/technology reducing compromise/breach likelihood.': 'Benefit placeholder 5',
        'Reduce employee burden/workload with automation, orchestration and AI': 'Benefit placeholder 6',
        'Risks assess supply chain compromise, data provenance, use cases threats.': 'Benefit placeholder 7',
        'Gap and improvement area Identification aligned with your business vision.': 'Benefit placeholder 8',
        'Upskilled team effectively using AI security tools to increase effectiveness': 'Benefit placeholder 9',
        'Tailored advice and guidance on relevant AI security updates': 'Benefit placeholder 10',
        
        # Service Definition subsections
        'AI Security advisory': 'Service Definition Subsection 1',
        'Secure Design and Development for AI systems': 'Service Definition Subsection 2',
        'Leverage AI to Protect from Cyber Attack': 'Service Definition Subsection 3',
        
        # Additional AI Security content in description and service definition
        'Use of AI is rapidly growing': '[Generic service description text]',
        'generative AI services such as ChatGPT': '[Generic service description text]',
        "organisation's use of AI": '[Generic service description text]',
        "'shadow' AI use": '[Generic service description text]',
        "shadow' AI use": '[Generic service description text]',
        'shadow AI use': '[Generic service description text]',
        "identifying 'shadow' AI": '[Generic service description text]',
        'identifying shadow AI': '[Generic service description text]',
        'develop innovative AI systems': '[Generic service description text]',
        'AI technology and threat landscapes': '[Generic service description text]',
        'distinct security challenges': '[Generic service description text]',
        'Security by Design principles': '[Generic service description text]',
        'leverage AI to do business': '[Generic service description text]',
        'leverage AI to enhance': '[Generic service description text]',
        'AI to enhance and strength': '[Generic service description text]',
        'AI to reduce': '[Generic service description text]',
        'AI systems': '[service systems]',
        'AI infrastructure': '[service infrastructure]',
        'AI capabilities': '[service capabilities]',
    }
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
    
    # Replace in runs (for formatted text)
    for para in doc.paragraphs:
        for run in para.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, new_text)
                    for run in para.runs:
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
    
    # Save the fixed template
    doc.save(str(template_path))
    print(f"Template fixed and saved: {template_path}")
    print(f"All AI Security content replaced with generic placeholders")
    
    return template_path

if __name__ == "__main__":
    fix_template()

