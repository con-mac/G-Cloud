"""
URGENT: Completely remove ALL AI Security content from template
Replace Service Definition sections with generic placeholders
"""

from pathlib import Path
from docx import Document
from datetime import datetime

def completely_fix_template():
    """Completely remove ALL AI Security content"""
    
    template_path = Path(__file__).parent.parent / "templates" / "service_description_template.docx"
    
    print(f"Loading template: {template_path}")
    
    if not template_path.exists():
        print(f"ERROR: Template not found at {template_path}")
        return
    
    doc = Document(str(template_path))
    
    print("Removing ALL AI Security content from template...")
    print()
    
    # Complete replacements - remove ALL AI Security content
    replacements = {
        # Title
        'AI Security': 'ENTER SERVICE NAME HERE',
        'AI security': 'enter service name here',
        
        # Description section
        'AI Security Advisory': '[SERVICE NAME]',
        'enables teams to understand and use AI ethically and securely': '[Generic service description placeholder text]',
        'Secure Design and Development of AI helps': '[Generic service description placeholder text]',
        'AI products are secure-by-design': '[Generic service description placeholder text]',
        'use case-relevant': '[Generic service description placeholder text]',
        
        # Service Definition subsections - COMPLETE replacement
        'ENTER SERVICE NAME HERE advisory': 'Service Definition Subsection 1',
        'AI Security advisory': 'Service Definition Subsection 1',
        
        # Remove ALL AI Security context from Service Definition content
        '[Generic service description text], often in unplanned and ungoverned ways, for example the huge uptake of [Generic service description text]': '[Generic service description placeholder text for service definition subsection]',
        'We can help security leaders understand and secure their [Generic service description text] – including identifying \'shadow\' AI use': '[Generic service description placeholder text for service definition subsection]',
        'We can help security leaders understand and secure their': '[Generic service description placeholder text]',
        'security leaders understand and secure': '[Generic service description placeholder text]',
        'identifying \'shadow\'': '[Generic service description text]',
        'identifying shadow': '[Generic service description text]',
        'including identifying \'shadow\'': '[Generic service description placeholder text]',
        'including identifying shadow': '[Generic service description placeholder text]',
        '– including identifying': '[Generic service description placeholder text]',
        'including identifying': '[Generic service description placeholder text]',
        'to address the unique challenges and threats associated with AI': '[Generic service description placeholder text]',
        'Secure Design and Development for AI systems': 'Service Definition Subsection 2',
        'We will work with you to develop innovative AI systems': '[Generic service description placeholder text]',
        'The nature of these – both their innovative business uses, and the evolving AI technology': '[Generic service description placeholder text]',
        'Leverage AI to Protect from Cyber Attack': 'Service Definition Subsection 3',
        'leverage AI to do business': '[Generic service description placeholder text]',
        'leverage AI to enhance': '[Generic service description placeholder text]',
        'Leverage [Generic service description text] your security programme': '[Generic service description placeholder text]',
        
        # Additional AI Security references
        'shadow AI use': '[Generic service description text]',
        'shadow\' AI use': '[Generic service description text]',
        "'shadow' AI use": '[Generic service description text]',
        'identifying \'shadow\' AI': '[Generic service description text]',
        'ChatGPT': '[Generic service description text]',
        'generative AI services': '[Generic service description text]',
        'AI systems': '[service systems]',
        'AI infrastructure': '[service infrastructure]',
        'AI capabilities': '[service capabilities]',
        'AI technology': '[service technology]',
        'AI use': '[service use]',
        'AI products': '[service products]',
        'AI-aware': '[service-aware]',
        'AI literacy': '[service area] literacy',
        'associated with AI': '[Generic service description text]',
        'AI defences': '[service defences]',
        'AI security': '[service security]',
    }
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
    
    # Replace in runs (for formatted text)
    for para in doc.paragraphs:
        for run in para.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, new_text)
                    for run in para.runs:
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
    
    # Save the fixed template
    doc.save(str(template_path))
    print(f"✅ Template fixed and saved: {template_path}")
    print(f"✅ All AI Security content removed")
    
    # Verify
    print()
    print("Verifying template is clean...")
    doc2 = Document(str(template_path))
    ai_security_found = []
    for para in doc2.paragraphs:
        for run in para.runs:
            text = run.text
            if text and ('AI Security' in text or 'ChatGPT' in text or ('AI' in text and 'shadow' in text.lower())):
                ai_security_found.append(text[:200])
    
    if ai_security_found:
        print(f"⚠️  Still found {len(ai_security_found)} AI Security references:")
        for text in ai_security_found[:5]:
            print(f"  - {text[:150]}")
    else:
        print("✅ Template is clean - no AI Security content found")
    
    return template_path

if __name__ == "__main__":
    completely_fix_template()

